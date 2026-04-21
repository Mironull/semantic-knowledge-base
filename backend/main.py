import sqlite3
import os
import io
from datetime import datetime
from typing import List
from uuid import uuid4

from fastapi import FastAPI, File, UploadFile, HTTPException, Query
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pydantic import BaseModel

# Import libraries for document parsing
try:
    from PyPDF2 import PdfReader
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# --- SQLite setup ---
DB_FILE = "docstore.db"

def get_db():
    """Return a connection to the SQLite database."""
    conn = sqlite3.connect(DB_FILE, detect_types=sqlite3.PARSE_DECLTYPES)
    conn.row_factory = sqlite3.Row  # to access columns by name
    return conn

def init_db():
    """Create the documents table if it doesn't exist."""
    with get_db() as conn:
        conn.execute("""
            CREATE TABLE IF NOT EXISTS documents (
                id TEXT PRIMARY KEY,
                filename TEXT NOT NULL,
                content_type TEXT NOT NULL,
                data BLOB NOT NULL,
                upload_date TIMESTAMP NOT NULL
            )
        """)
        conn.execute("CREATE INDEX IF NOT EXISTS idx_filename ON documents(filename)")

# Initialize the database when the app starts
init_db()

app = FastAPI(title="Document Store API (SQLite)", version="1.0.0")

# Add CORS middleware to allow frontend requests
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000"],  # React dev server
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

class DocumentMetadata(BaseModel):
    id: str
    filename: str
    content_type: str
    upload_date: datetime

class DocumentMetadataWithSize(BaseModel):
    id: str
    filename: str
    content_type: str
    upload_date: datetime
    size: int  # size in bytes

@app.post("/upload", response_model=DocumentMetadata)
async def upload_document(file: UploadFile = File(...)):
    """
    Upload a document and store it in SQLite.
    Returns metadata of the stored document.
    """
    content = await file.read()
    doc_id = str(uuid4())
    now = datetime.utcnow()

    with get_db() as conn:
        conn.execute(
            "INSERT INTO documents (id, filename, content_type, data, upload_date) VALUES (?, ?, ?, ?, ?)",
            (doc_id, file.filename, file.content_type, content, now)
        )

    return DocumentMetadata(id=doc_id, filename=file.filename, content_type=file.content_type, upload_date=now)

@app.get("/download/{doc_id}")
async def download_document(doc_id: str):
    """
    Download a document by its ID.
    """
    with get_db() as conn:
        row = conn.execute("SELECT filename, content_type, data FROM documents WHERE id = ?", (doc_id,)).fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="Document not found")

    return StreamingResponse(
        iter([row["data"]]),
        media_type=row["content_type"],
        headers={"Content-Disposition": f"attachment; filename={row['filename']}"}
    )

@app.get("/search", response_model=List[DocumentMetadata])
async def search_documents(name: str = Query(..., min_length=1)):
    """
    Search for documents by filename (case‑insensitive partial match).
    """
    with get_db() as conn:
        rows = conn.execute(
            "SELECT id, filename, content_type, upload_date FROM documents WHERE filename LIKE ?",
            (f"%{name}%",)  # SQLite LIKE is case‑insensitive by default only for ASCII; use COLLATE NOCASE if needed
        ).fetchall()

    return [
        DocumentMetadata(
            id=row["id"],
            filename=row["filename"],
            content_type=row["content_type"],
            upload_date=datetime.fromisoformat(row["upload_date"]) if isinstance(row["upload_date"], str) else row["upload_date"]
        )
        for row in rows
    ]

@app.get("/documents", response_model=List[DocumentMetadataWithSize])
async def list_all_documents():
    """
    List all documents in the database with metadata including file size.
    """
    with get_db() as conn:
        rows = conn.execute(
            "SELECT id, filename, content_type, upload_date, length(data) as size FROM documents ORDER BY upload_date DESC"
        ).fetchall()

    return [
        DocumentMetadataWithSize(
            id=row["id"],
            filename=row["filename"],
            content_type=row["content_type"],
            upload_date=datetime.fromisoformat(row["upload_date"]) if isinstance(row["upload_date"], str) else row["upload_date"],
            size=row["size"]
        )
        for row in rows
    ]

@app.get("/preview/{doc_id}")
async def preview_document(doc_id: str):
    """
    Preview document content. Returns text content for supported formats including PDF and DOCX.
    """
    with get_db() as conn:
        row = conn.execute("SELECT filename, content_type, data FROM documents WHERE id = ?", (doc_id,)).fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="Document not found")

    filename = row["filename"]
    content_type = row["content_type"] or ""
    data = row["data"]

    # Try to decode as text for text-based formats
    if content_type and (
        "text" in content_type.lower() or
        "json" in content_type.lower() or
        "xml" in content_type.lower() or
        "javascript" in content_type.lower() or
        content_type in ["application/json", "application/xml"]
    ):
        try:
            text_content = data.decode('utf-8')
            return {"content": text_content, "type": "text", "content_type": content_type}
        except UnicodeDecodeError:
            pass

    # Handle PDF files
    if PDF_AVAILABLE and (
        "pdf" in content_type.lower() or
        filename.lower().endswith('.pdf')
    ):
        try:
            pdf_file = io.BytesIO(data)
            pdf_reader = PdfReader(pdf_file)
            text_content = []

            for page_num, page in enumerate(pdf_reader.pages, 1):
                page_text = page.extract_text()
                if page_text.strip():
                    text_content.append(f"=== Страница {page_num} ===\n{page_text}\n")

            if text_content:
                return {
                    "content": "\n".join(text_content),
                    "type": "text",
                    "content_type": content_type,
                    "pages": len(pdf_reader.pages)
                }
            else:
                return {
                    "content": "PDF файл не содержит текста или текст не может быть извлечен",
                    "type": "unsupported",
                    "content_type": content_type
                }
        except Exception as e:
            return {
                "content": f"Ошибка при чтении PDF: {str(e)}",
                "type": "error",
                "content_type": content_type
            }

    # Handle DOCX files
    if DOCX_AVAILABLE and (
        "wordprocessingml" in content_type.lower() or
        "msword" in content_type.lower() or
        filename.lower().endswith('.docx')
    ):
        try:
            docx_file = io.BytesIO(data)
            doc = Document(docx_file)
            text_content = []

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_content.append(row_text)

            if text_content:
                return {
                    "content": "\n\n".join(text_content),
                    "type": "text",
                    "content_type": content_type
                }
            else:
                return {
                    "content": "DOCX файл не содержит текста",
                    "type": "unsupported",
                    "content_type": content_type
                }
        except Exception as e:
            return {
                "content": f"Ошибка при чтении DOCX: {str(e)}",
                "type": "error",
                "content_type": content_type
            }

    # For other binary formats, return metadata only
    return {
        "content": f"Предпросмотр недоступен для формата {content_type or 'неизвестный формат'}",
        "type": "unsupported",
        "content_type": content_type,
        "size": len(data)
    }

@app.get("/")
async def root():
    return {"message": "Document Store API (SQLite) is running. Use /docs for Swagger UI."}