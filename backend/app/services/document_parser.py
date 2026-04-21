"""
Document parser for extracting text from various file formats.
"""
import io
from abc import ABC, abstractmethod
from typing import Dict, Any

# Import libraries for document parsing
try:
    from PyPDF2 import PdfReader
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class BaseParser(ABC):
    """Abstract base class for document parsers."""

    @abstractmethod
    def can_parse(self, filename: str, content_type: str) -> bool:
        """
        Check if this parser can handle the given file.

        Args:
            filename: Name of the file
            content_type: MIME type of the file

        Returns:
            bool: True if parser can handle this file
        """
        pass

    @abstractmethod
    def parse(self, data: bytes) -> Dict[str, Any]:
        """
        Parse document and extract text content.

        Args:
            data: Binary file data

        Returns:
            Dict containing content, type, and optional metadata
        """
        pass


class TextParser(BaseParser):
    """Parser for plain text files (txt, json, xml, etc.)."""

    def can_parse(self, filename: str, content_type: str) -> bool:
        """Check if file is a text-based format."""
        return any([
            "text" in content_type.lower(),
            "json" in content_type.lower(),
            "xml" in content_type.lower(),
            "javascript" in content_type.lower(),
            content_type in ["application/json", "application/xml"]
        ])

    def parse(self, data: bytes) -> Dict[str, Any]:
        """Extract text from text file."""
        try:
            text_content = data.decode('utf-8')
            return {
                "content": text_content,
                "type": "text"
            }
        except UnicodeDecodeError:
            return {
                "content": "Не удалось декодировать файл как текст",
                "type": "error"
            }


class PDFParser(BaseParser):
    """Parser for PDF documents."""

    def can_parse(self, filename: str, content_type: str) -> bool:
        """Check if file is a PDF."""
        if not PDF_AVAILABLE:
            return False
        return "pdf" in content_type.lower() or filename.lower().endswith('.pdf')

    def parse(self, data: bytes) -> Dict[str, Any]:
        """Extract text from PDF file."""
        try:
            pdf_file = io.BytesIO(data)
            pdf_reader = PdfReader(pdf_file)
            text_content = []

            for page_num, page in enumerate(pdf_reader.pages, 1):
                page_text = page.extract_text()
                if page_text.strip():
                    text_content.append(f"=== Страница {page_num} ===\n{page_text}\n")

            if text_content:
                return {
                    "content": "\n".join(text_content),
                    "type": "text",
                    "pages": len(pdf_reader.pages)
                }
            else:
                return {
                    "content": "PDF файл не содержит текста или текст не может быть извлечен",
                    "type": "unsupported"
                }
        except Exception as e:
            return {
                "content": f"Ошибка при чтении PDF: {str(e)}",
                "type": "error"
            }


class DOCXParser(BaseParser):
    """Parser for Microsoft Word documents."""

    def can_parse(self, filename: str, content_type: str) -> bool:
        """Check if file is a DOCX."""
        if not DOCX_AVAILABLE:
            return False
        return any([
            "wordprocessingml" in content_type.lower(),
            "msword" in content_type.lower(),
            filename.lower().endswith('.docx')
        ])

    def parse(self, data: bytes) -> Dict[str, Any]:
        """Extract text from DOCX file."""
        try:
            docx_file = io.BytesIO(data)
            doc = Document(docx_file)
            text_content = []

            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    )
                    if row_text:
                        text_content.append(row_text)

            if text_content:
                return {
                    "content": "\n\n".join(text_content),
                    "type": "text"
                }
            else:
                return {
                    "content": "DOCX файл не содержит текста",
                    "type": "unsupported"
                }
        except Exception as e:
            return {
                "content": f"Ошибка при чтении DOCX: {str(e)}",
                "type": "error"
            }


class DocumentParserService:
    """Service for parsing documents using multiple parsers."""

    def __init__(self):
        """Initialize parser service with available parsers."""
        self.parsers = [
            TextParser(),
            PDFParser(),
            DOCXParser()
        ]

    def parse_document(
        self,
        data: bytes,
        filename: str,
        content_type: str
    ) -> Dict[str, Any]:
        """
        Parse document using appropriate parser.

        Args:
            data: Binary file data
            filename: Name of the file
            content_type: MIME type of the file

        Returns:
            Dict containing parsed content and metadata
        """
        # Find appropriate parser
        for parser in self.parsers:
            if parser.can_parse(filename, content_type):
                result = parser.parse(data)
                result["content_type"] = content_type
                return result

        # No parser found
        return {
            "content": f"Предпросмотр недоступен для формата {content_type or 'неизвестный формат'}",
            "type": "unsupported",
            "content_type": content_type,
            "size": len(data)
        }
